import pandas as pd
from docx import Document

# Excel file padhein
df = pd.read_excel('Sample.xlsx')  # Apni Excel file ka naam

# Star rating function (sirf .0 aur .5 ke liye)
def number_to_stars(rating):
    full_stars = int(rating)  # Full stars
    half_star = (rating - full_stars) == 0.5  # Half star check
    empty_stars = 5 - full_stars - (1 if half_star else 0)
    
    stars = '★' * full_stars  # Full stars
    if half_star:
        stars += '⯪'  # Half star
    stars += '☆' * empty_stars  # Empty stars
    return stars

# Word document banayein
doc = Document()
for index, row in df.iterrows():
    doc.add_paragraph(f"Roll Number: {row['Roll Number']}")
    doc.add_paragraph(f"Name: {row['Name']}")
    doc.add_paragraph(f"Email ID: {row['Email ID']}")
    doc.add_paragraph(f"Citi Score: {row['Citi Score']}")
    doc.add_paragraph(f"Rating: {number_to_stars(row['Citi Score'])}")
    doc.add_paragraph("------------------------")
    # Har row ke baad naya page start karo
    if index < len(df) - 1:  # Last row ke liye page break na add karein
        doc.add_page_break()

# File save kijiye
doc.save('new page.docx')