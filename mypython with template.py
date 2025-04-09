import pandas as pd
from docx import Document
from copy import deepcopy
from docx.enum.text import WD_BREAK  # WD_BREAK import kiya gaya

# Excel file padhein
df = pd.read_excel('Sample.xlsx')

# Star rating function (sirf .0 aur .5 ke liye)
def number_to_stars(rating):
    full_stars = int(rating)
    half_star = (rating - full_stars) == 0.5
    empty_stars = 5 - full_stars - (1 if half_star else 0)
    
    stars = '★' * full_stars
    if half_star:
        stars += '⯪'
    stars += '☆' * empty_stars
    return stars

# Template file load karein
template = Document('template.docx')  # Apna template file ka naam yahan daaliye

# Naya document banayein jisme sab data jayega
doc = Document()

for index, row in df.iterrows():
    # Har row ke liye template ki deep copy banayein
    current_doc = deepcopy(template)
    
    # Har paragraph mein placeholders replace karein
    for paragraph in current_doc.paragraphs:
        if '{{Roll Number}}' in paragraph.text:
            paragraph.text = paragraph.text.replace('{{Roll Number}}', str(row['Roll Number']))
        if '{{Name}}' in paragraph.text:
            paragraph.text = paragraph.text.replace('{{Name}}', str(row['Name']))
        if '{{Email ID}}' in paragraph.text:
            paragraph.text = paragraph.text.replace('{{Email ID}}', str(row['Email ID']))
        if '{{Citi Score}}' in paragraph.text:
            paragraph.text = paragraph.text.replace('{{Citi Score}}', str(row['Citi Score']))
        if '{{Rating}}' in paragraph.text:
            paragraph.text = paragraph.text.replace('{{Rating}}', number_to_stars(row['Citi Score']))

    # Current document ke paragraphs ko main document mein add karein
    for paragraph in current_doc.paragraphs:
        new_para = doc.add_paragraph(paragraph.text)  # Naya paragraph add karein
        # Paragraph ke formatting ko copy karein (agar hai)
        if paragraph.paragraph_format.space_before:
            new_para.paragraph_format.space_before = paragraph.paragraph_format.space_before
        if paragraph.paragraph_format.space_after:
            new_para.paragraph_format.space_after = paragraph.paragraph_format.space_after

    # Har row ke baad page break add karein (last row ke liye nahi)
    if index < len(df) - 1:
        doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

# File save kijiye
doc.save('output dges.docx')